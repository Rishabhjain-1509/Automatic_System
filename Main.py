# -*- coding: utf-8 -*-
"""
Created on Thu May 25 10:07:00 2020

@author: Rishabh Jain
"""

# Import modules
#Library for document coversion
import docx  
#Library for extract the text from image.                            
import pytesseract 
#Library use for sreaching the text in image.                     
from PIL import Image                    
#library for Some attribute of the document (Style , Spacing, etc....)
from docx.shared import Inches  
#Library for check the Spelling in the give pragharph.         
from spellchecker import SpellChecker


                    
# Install the Pytesseract in our project so its easy to convert image to text.
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\Rishabh Jain\AppData\Local\Tesseract-OCR\tesseract.exe"

# Create the image directy to find the text.
image = Image.open('image.png')

# Pass image in Pytesseract and sreach for the text in the image.
image_to_text = pytesseract.image_to_string(image, lang='eng')

# Print the text in the image.
print(image_to_text)

# Split all the text and store in the array to check the spelling and correct them.
data = image_to_text.split()

#print the array of split.
print(data)

 # Add the function which check the spelling in the array.
spell = SpellChecker()
 
# Help to find the unknown word in the array.
misspelled = spell.unknown(data)
 
# Loop for check each word in the Array.
for word in misspelled:

    print(spell.correction(word))
 
    print(spell.candidates(word))   


doc = docx.Document() 
  
doc_para = doc.add_paragraph(image_to_text) 

#doc_para.add_run().italic = True


doc.save('A:\Study\Rishabh Jain\SPIT\SY_SEM-4\Mini Project\Sample.txt')