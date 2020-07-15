# -*- coding: utf-8 -*-
"""
Created on Sat Apr 18 09:42:47 2020

@author: Rishabh Jain
"""


# Import modules
from PIL import Image
import pytesseract


# Include tesseract executable in your path
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\Rishabh Jain\AppData\Local\Tesseract-OCR\tesseract.exe"

# Create an image object of PIL library
image = Image.open('image.png')

# pass image into pytesseract module
# pytesseract is trained in many languages
image_to_text = pytesseract.image_to_string(image, lang='eng')

# Print the text
print(image_to_text)


# import docx NOT python-docx 
import docx
from docx.shared import Inches 
  
# create an instance of a word document 
doc = docx.Document() 
  
# add a heading of level 0 (largest heading) 
doc.add_heading('Heading for the document', 0) 
  
# add a paragraph and store  
# the object in a variable 
doc_para = doc.add_paragraph('Your paragraph goes here, ') 
  
# add a run i.e, style like  
# bold, italic, underline, etc. 
doc_para.add_run(image_to_text).bold = True
doc_para.add_run(', and ') 
doc_para.add_run('these words are italic').italic = True
  
# add a page break to start a new page 
doc.add_page_break() 
  
# add a heading of level 2 
doc.add_heading('Heading level 2', 2) 
  
# pictures can also be added to our word document 
# width is optional 
#doc.add_picture('path_to_picture') 
  
# now save the document to a location 
doc.save('A:\Mini Project\Sample.txt')